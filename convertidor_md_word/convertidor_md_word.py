"""Markdown to Word converter for FinanceFlow documents.

Usage:
    python convertidor_md_word.py input.md output.docx

The converter handles the most common Markdown structures used in the
FinanceFlow documentation:
- headings
- paragraphs
- bullet and numbered lists
- fenced code blocks
- tables
- images
- blockquotes
- horizontal rules

It depends on python-docx.
"""

from __future__ import annotations

import argparse
import os
import re
from dataclasses import dataclass
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_PATTERN = re.compile(r"^(\s*)([-*+])\s+(.*)$")
NUMBERED_PATTERN = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
QUOTE_PATTERN = re.compile(r"^>\s?(.*)$")
TABLE_SEPARATOR_PATTERN = re.compile(r"^\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?$")
IMAGE_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


@dataclass
class ListContext:
    kind: str
    level: int


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")
        borders.append(element)
    tbl_pr.append(borders)


def apply_document_style(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"):
        if name in styles:
            styles[name].font.name = "Calibri"


def is_table_divider(line: str) -> bool:
    return bool(TABLE_SEPARATOR_PATTERN.match(line.strip()))


def split_table_row(line: str) -> List[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text: str) -> None:
    cursor = 0
    for match in IMAGE_PATTERN.finditer(text):
        before = text[cursor:match.start()]
        if before:
            add_linked_text(paragraph, before)
        alt_text, image_path = match.groups()
        if os.path.exists(image_path):
            paragraph.add_run().add_picture(image_path, width=Inches(5.8))
        else:
            paragraph.add_run(f"[Imagen: {alt_text}] {image_path}")
        cursor = match.end()
    remaining = text[cursor:]
    if remaining:
        add_linked_text(paragraph, remaining)


def add_linked_text(paragraph, text: str) -> None:
    cursor = 0
    for match in LINK_PATTERN.finditer(text):
        before = text[cursor:match.start()]
        if before:
            paragraph.add_run(before)
        label, url = match.groups()
        add_hyperlink(paragraph, label, url)
        cursor = match.end()
    tail = text[cursor:]
    if tail:
        paragraph.add_run(tail)


def add_paragraph(document: Document, text: str, style: Optional[str] = None):
    paragraph = document.add_paragraph(style=style)
    add_inline_runs(paragraph, text)
    return paragraph


def parse_table(document: Document, lines: List[str], start_index: int) -> int:
    header = split_table_row(lines[start_index])
    index = start_index + 1
    if index < len(lines) and is_table_divider(lines[index]):
        index += 1

    rows: List[List[str]] = []
    while index < len(lines):
        line = lines[index]
        if not line.strip().startswith("|"):
            break
        rows.append(split_table_row(line))
        index += 1

    if not header:
        return index - 1

    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for col_index, value in enumerate(header):
        hdr_cells[col_index].text = value
        set_cell_shading(hdr_cells[col_index], "D9EAF7")

    for row in rows:
        row_cells = table.add_row().cells
        for col_index in range(len(header)):
            row_cells[col_index].text = row[col_index] if col_index < len(row) else ""

    set_table_borders(table)
    return index - 1


def normalize_indentation(raw: str) -> int:
    return len(raw) - len(raw.lstrip(" "))


def markdown_to_docx(markdown_text: str, output_path: str) -> None:
    document = Document()
    apply_document_style(document)

    lines = markdown_text.splitlines()
    i = 0
    in_code_block = False
    code_lines: List[str] = []
    code_language = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                paragraph = document.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                paragraph.style = document.styles["No Spacing"] if "No Spacing" in document.styles else paragraph.style
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.25)
                paragraph_format.space_before = Pt(3)
                paragraph_format.space_after = Pt(3)
                in_code_block = False
                code_lines = []
                code_language = ""
            else:
                in_code_block = True
                code_language = stripped[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip("\n"))
            i += 1
            continue

        if not stripped:
            document.add_paragraph("")
            i += 1
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            document.add_heading(text, level=min(level, 6))
            i += 1
            continue

        if stripped == "---" or stripped == "***":
            paragraph = document.add_paragraph()
            paragraph.add_run("_" * 48)
            i += 1
            continue

        quote_match = QUOTE_PATTERN.match(stripped)
        if quote_match:
            paragraph = document.add_paragraph(style="Intense Quote" if "Intense Quote" in document.styles else None)
            add_inline_runs(paragraph, quote_match.group(1))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_divider(lines[i + 1]):
            i = parse_table(document, lines, i) + 1
            continue

        bullet_match = BULLET_PATTERN.match(line)
        numbered_match = NUMBERED_PATTERN.match(line)
        if bullet_match or numbered_match:
            match = bullet_match or numbered_match
            assert match is not None
            indent = normalize_indentation(match.group(1))
            level = indent // 2
            style = "List Bullet" if bullet_match else "List Number"
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            add_inline_runs(paragraph, match.group(3).strip())
            i += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, stripped)
        i += 1

    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown files to Word (.docx) documents.")
    parser.add_argument("input", help="Path to the input Markdown file")
    parser.add_argument("output", help="Path to the output Word document (.docx)")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as handle:
        markdown_text = handle.read()

    markdown_to_docx(markdown_text, args.output)
    print(f"Created {args.output}")


if __name__ == "__main__":
    main()
